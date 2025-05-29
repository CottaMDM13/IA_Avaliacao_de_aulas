import json
import os
import sqlite3
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configurar logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def generate_final_report(video_results, audio_results, text_results, metrics, output_dir, video_name):
    try:
        logger.info(f"Gerando relatório para: {video_name}")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report = {
            "video": video_results,
            "audio": audio_results,
            "text": text_results,
            "overall_score": metrics["overall_score"],
            "approved": metrics["approved"],
            "feedback": metrics["feedback"],
            "timestamp": timestamp,
            "video_name": video_name
        }
        os.makedirs(output_dir, exist_ok=True)
        conn = sqlite3.connect(f"{output_dir}/reports.db")
        cursor = conn.cursor()
        
        # Criar tabela se não existir
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS reports (
                id INTEGER PRIMARY KEY,
                video_name TEXT,
                score REAL,
                feedback TEXT,
                approved BOOLEAN,
                timestamp TEXT
            )
        """)
        
        # Verificar se a coluna 'approved' existe
        cursor.execute("PRAGMA table_info(reports)")
        columns = [col[1] for col in cursor.fetchall()]
        if 'approved' not in columns:
            cursor.execute("ALTER TABLE reports ADD COLUMN approved BOOLEAN")
        
        cursor.execute("""
            INSERT INTO reports (video_name, score, feedback, approved, timestamp)
            VALUES (?, ?, ?, ?, ?)
        """, (video_name, metrics["overall_score"], json.dumps(metrics["feedback"]), metrics["approved"], timestamp))
        conn.commit()
        conn.close()
        
        base_name = f"final_report_{video_name.rsplit('.', 1)[0]}_{timestamp}"
        json_path = f"{output_dir}/{base_name}.json"
        txt_path = f"{output_dir}/{base_name}.txt"
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(report, f, indent=4, ensure_ascii=False)
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(f"Relatório Final - {video_name} - {timestamp}\n")
            f.write(f"Score Geral: {metrics['overall_score']*100:.0f}%\n")
            f.write(f"Status: {'APROVADO' if metrics['approved'] else 'REPROVADO'}\n\n")
            f.write("Resumo:\n")
            f.write(f"{text_results.get('review', 'Sem resumo disponível.')}\n\n")
            f.write("Feedback Técnico:\n")
            for fb in metrics["feedback"]:
                f.write(f"- {fb['aspect']}: {fb['comment']} (Nota: {fb['score']*100:.0f}%)\n")
                if fb["suggestion"]:
                    f.write(f"  Sugestão: {fb['suggestion']}\n")
        logger.info(f"Relatório gerado: {json_path}")
        return {"json_path": json_path, "txt_path": txt_path, "metrics": report}
    except Exception as e:
        logger.error(f"Erro ao gerar relatório para {video_name}: {str(e)}")
        raise

def generate_docx_report(video_name, metrics, output_dir):
    try:
        logger.debug(f"Gerando .docx para: {video_name}")
        doc = Document()
        
        # Cabeçalho
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "Relatório de Avaliação de Videoaulas"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.runs[0].font.size = Pt(12)
        
        # Rodapé com número da página
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer_para.add_run("Página ")
        footer_run.font.size = Pt(10)
        
        # Adicionar campo de número da página
        fld_simple = OxmlElement('w:fldSimple')
        fld_simple.set(qn('w:instr'), 'PAGE')
        run = OxmlElement('w:r')
        run.append(OxmlElement('w:t'))
        fld_simple.append(run)
        footer_para._p.append(fld_simple)
        
        # Título
        doc.add_heading(f'Relatório de Avaliação - {video_name}', 0)
        doc.add_paragraph(f'Data/Hora: {metrics["timestamp"]}').style.font.size = Pt(12)
        doc.add_paragraph(f'Score Geral: {metrics["overall_score"]*100:.0f}%').style.font.size = Pt(12)
        doc.add_paragraph(f'Status: {"APROVADO" if metrics["approved"] else "REPROVADO"}').style.font.size = Pt(12)
        
        doc.add_heading('Feedback Detalhado', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Aspecto'
        hdr_cells[1].text = 'Nota (0-100)'
        hdr_cells[2].text = 'Comentário'
        hdr_cells[3].text = 'Sugestões de Melhoria'
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        for fb in metrics["feedback"]:
            row_cells = table.add_row().cells
            row_cells[0].text = fb["aspect"]
            row_cells[1].text = f'{fb["score"]*100:.0f}'
            row_cells[2].text = fb["comment"]
            row_cells[3].text = fb["suggestion"] if fb["suggestion"] else 'Satisfatório, continue assim!'
            
            score = fb["score"]
            color = RGBColor(40, 167, 69) if score >= 0.9 else RGBColor(255, 193, 7) if score >= 0.6 else RGBColor(220, 53, 69)
            row_cells[1].paragraphs[0].runs[0].font.color.rgb = color
        
        base_name = f"report_{video_name.rsplit('.', 1)[0]}_{metrics['timestamp'].replace('/', '').replace(':', '').replace(' ', '_')}"
        docx_path = f"{output_dir}/{base_name}.docx"
        doc.save(docx_path)
        logger.info(f"Arquivo .docx gerado: {docx_path}")
        return docx_path
    except Exception as e:
        logger.error(f"Erro ao gerar .docx para {video_name}: {str(e)}")
        raise
